#!/usr/bin/env python3
"""
Generate the submission-ready DOCX of the publication.

Prose and tables come from scripts/publication_content.py; figures come from
scripts/make_figures.py. This script only lays them out in Word format.

Previously this file held its own copy of the entire paper, including a
hardcoded comparison table of numbers that had never been measured, and its own
copy of the figure code. Both now have a single home, so the DOCX and the LaTeX
version cannot disagree.

Usage:
    python scripts/generate_publication_docx.py
    python scripts/generate_publication_docx.py --trainer-state runs/main_full/trainer_state.json
    python scripts/generate_publication_docx.py --skip-figures
"""

from __future__ import annotations

import argparse
import sys
from pathlib import Path

SCRIPT_DIR = Path(__file__).resolve().parent
PROJECT_ROOT = SCRIPT_DIR.parent
for d in (PROJECT_ROOT, SCRIPT_DIR):
    if str(d) not in sys.path:
        sys.path.insert(0, str(d))

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Install python-docx: pip install python-docx", file=sys.stderr)
    sys.exit(1)

import publication_content as P

DEFAULT_FIG_DIR = PROJECT_ROOT / "publication" / "figures"
DEFAULT_OUTPUT_DOCX = (PROJECT_ROOT / "publication"
                       / "Haque_etal_LLM_Bangladesh_Labour_Law.docx")

# (filename, caption) in the order they appear in the paper.
FIGURE_CAPTIONS = {
    "figure1_dataset_pipeline.png": "Figure 1: Dataset creation and validation pipeline.",
    "figure2_finetuning_pipeline.png": "Figure 2: Fine-tuning and deployment pipeline.",
    "figure3_evaluation_design.png": "Figure 3: Comparative evaluation design.",
    "figure4_training_loss.png": "Figure 4: Training and validation loss.",
    "figure5_eval_loss_perplexity.png": "Figure 5: Validation loss and perplexity.",
    "figure6_dataset_statistics.png": "Figure 6: Dataset statistics.",
    "figure7_results_comparison.png": "Figure 7: System comparison.",
    "figure8_error_analysis.png": "Figure 8: Failure modes by system.",
}


def txt(block: str) -> list[str]:
    """Prose block -> plain-text paragraphs with numeric citations."""
    return P.paragraphs(P.strip_latex(P.render_citations(block, "plain")))


def add_body(doc, block):
    for para in txt(block):
        doc.add_paragraph(para)


def add_caption(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(9)
        run.italic = True


def add_figure(doc, fig_dir, name, width=6.0):
    path = Path(fig_dir) / name
    if not path.exists():
        return False
    doc.add_picture(str(path), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_caption(doc, FIGURE_CAPTIONS.get(name, name))
    return True


def add_table(doc, spec, number=None):
    """Render a (caption, header, rows, note) tuple as a Word table."""
    caption, header, rows, note = spec
    cap = P.strip_latex(caption or "")
    if number:
        cap = f"Table {number}: {cap}"
    if not rows:
        p = doc.add_paragraph(P.strip_latex(note or P.NOT_MEASURED))
        for run in p.runs:
            run.italic = True
        return False
    add_caption(doc, cap)
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = str(h)
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    for r in rows:
        cells = table.add_row().cells
        for i, v in enumerate(r):
            cells[i].text = str(v)
            for p in cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph()
    return True


def build_document(fig_dir, trainer_state_path=None) -> Document:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    # ---- title block ----
    t = doc.add_paragraph(P.TITLE)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in t.runs:
        run.bold = True
        run.font.size = Pt(16)
    a = doc.add_paragraph("\n".join(P.AUTHORS))
    a.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("Abstract", level=1)
    add_body(doc, P.ABSTRACT)

    doc.add_heading("1. Introduction", level=1)
    add_body(doc, P.INTRO)

    doc.add_heading("2. Related Work", level=1)
    add_body(doc, P.RELATED)

    doc.add_heading("3. Methodology", level=1)
    doc.add_heading("3.1 Dataset creation and validation", level=2)
    add_body(doc, P.METHOD_DATASET)
    add_figure(doc, fig_dir, "figure1_dataset_pipeline.png", 4.6)
    tno = 1
    if add_table(doc, P.table_dataset(), tno):
        tno += 1
    add_figure(doc, fig_dir, "figure6_dataset_statistics.png")

    doc.add_heading("3.2 Applied HR scenarios", level=2)
    add_body(doc, P.METHOD_SCENARIOS)
    if add_table(doc, P.table_scenario_verification(), tno):
        tno += 1
    examples = P.example_qa_pairs(3)
    if examples:
        doc.add_paragraph("Three representative items, one per topic:")
        for ex in examples:
            p = doc.add_paragraph()
            p.add_run(f"[{ex['topic']}] ").bold = True
            p.add_run(ex["question"])
            q = doc.add_paragraph()
            q.add_run("Reference: ").italic = True
            q.add_run(ex["reference"])
            q.add_run(f"  Gold sections: {', '.join(str(s) for s in ex['gold_sections'])}")

    doc.add_heading("3.3 Model and training", level=2)
    add_body(doc, P.METHOD_TRAINING)
    add_figure(doc, fig_dir, "figure2_finetuning_pipeline.png", 4.6)

    doc.add_heading("3.4 Deployment", level=2)
    add_body(doc, P.METHOD_DEPLOY)

    doc.add_heading("4. Experimental setup", level=1)
    add_body(doc, P.eval_setup_text())
    add_figure(doc, fig_dir, "figure3_evaluation_design.png")

    doc.add_heading("5. Results", level=1)
    doc.add_heading("5.1 Training", level=2)
    if add_table(doc, P.table_training(trainer_state_path), tno):
        tno += 1
    if add_table(doc, P.table_training_steps(trainer_state_path), tno):
        tno += 1
    add_figure(doc, fig_dir, "figure4_training_loss.png", 4.8)
    add_figure(doc, fig_dir, "figure5_eval_loss_perplexity.png")

    doc.add_heading("5.2 Answer quality", level=2)
    if add_table(doc, P.table_main_comparison("heldout"), tno):
        tno += 1
    if add_table(doc, P.table_main_comparison("scenario"), tno):
        tno += 1
    add_figure(doc, fig_dir, "figure7_results_comparison.png")

    doc.add_heading("5.3 Statistical comparison", level=2)
    if add_table(doc, P.table_significance("heldout"), tno):
        tno += 1
    if add_table(doc, P.table_significance("scenario"), tno):
        tno += 1

    doc.add_heading("5.4 Citation accuracy", level=2)
    if add_table(doc, P.table_citation("scenario"), tno):
        tno += 1

    doc.add_heading("5.5 Scope discipline", level=2)
    if add_table(doc, P.table_refusal(), tno):
        tno += 1

    doc.add_heading("5.6 Error analysis", level=2)
    if add_table(doc, P.table_error_analysis("heldout"), tno):
        tno += 1
    add_figure(doc, fig_dir, "figure8_error_analysis.png")

    doc.add_heading("5.7 Ablation", level=2)
    if add_table(doc, P.table_ablation(), tno):
        tno += 1
    add_body(doc, P.training_findings_text())

    doc.add_heading("6. Discussion", level=1)
    add_body(doc, P.discussion_text())

    doc.add_heading("7. Limitations", level=1)
    add_body(doc, P.LIMITATIONS)

    doc.add_heading("8. Conclusion", level=1)
    add_body(doc, P.CONCLUSION)

    doc.add_heading("Appendix: reproducibility", level=1)
    doc.add_paragraph(P.strip_latex(P.reproducibility_note()))

    doc.add_heading("References", level=1)
    for i, (_, text) in enumerate(P.REFERENCES, 1):
        doc.add_paragraph(f"[{i}] {P.strip_latex(text)}")

    return doc


def main():
    ap = argparse.ArgumentParser(description=__doc__,
                                 formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("--fig-dir", type=Path, default=DEFAULT_FIG_DIR)
    ap.add_argument("--output", type=Path, default=DEFAULT_OUTPUT_DOCX)
    ap.add_argument("--trainer-state", type=Path, default=None)
    ap.add_argument("--skip-figures", action="store_true",
                    help="reuse the PNGs already on disk")
    args = ap.parse_args()

    if not args.skip_figures:
        import make_figures
        make_figures.build_all(args.fig_dir, args.trainer_state)

    doc = build_document(args.fig_dir, args.trainer_state)
    args.output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(args.output))
    print(f"Wrote {args.output}")


if __name__ == "__main__":
    main()
